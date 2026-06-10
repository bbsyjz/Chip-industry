"""
芯片产业链每日情报 - DOCX报告生成器
"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime


STOCK_INFO = {
    "601216": {"name": "沪硅产业", "category": "上游材料", "segment": "硅片", "main_business": "半导体硅片研发、生产与销售", "position": "半导体制造的上游材料供应商"},
    "688126": {"name": "上海新昇", "category": "上游材料", "segment": "硅片", "main_business": "12英寸半导体硅片", "position": "国内12英寸硅片龙头"},
    "002163": {"name": "彩虹股份", "category": "上游材料", "segment": "玻璃基板", "main_business": "液晶基板玻璃、盖板玻璃", "position": "LCD面板上游材料"},
    "002236": {"name": "东旭光电", "category": "上游材料", "segment": "玻璃基板", "main_business": "光电显示器件玻璃基板", "position": "国内最大玻璃基板生产基地"},
    "600183": {"name": "生益科技", "category": "上游材料", "segment": "PCB基材", "main_business": "覆铜板、ABF薄膜材料", "position": "国内覆铜板龙头"},
    "603260": {"name": "合盛硅业", "category": "上游材料", "segment": "电子级多晶硅", "main_business": "工业硅、有机硅", "position": "硅基材料上游供应商"},
    "688221": {"name": "华大九天", "category": "上游设备", "segment": "EDA工具", "main_business": "EDA软件工具开发", "position": "国内EDA龙头"},
    "688036": {"name": "芯愿景", "category": "上游设备", "segment": "EDA工具", "main_business": "集成电路设计服务", "position": "IC分析设计服务商"},
    "688012": {"name": "中微公司", "category": "上游设备", "segment": "刻蚀设备", "main_business": "等离子刻蚀设备、MOCVD", "position": "半导体制造核心设备"},
    "002371": {"name": "北方华创", "category": "上游设备", "segment": "薄膜沉积/刻蚀设备", "main_business": "半导体设备", "position": "国内半导体设备龙头"},
    "688408": {"name": "华兴源创", "category": "上游设备", "segment": "封装测试设备", "main_business": "检测设备、自动化设备", "position": "面板检测龙头"},
    "300751": {"name": "迈为股份", "category": "上游设备", "segment": "封装测试设备", "main_business": "光伏电池丝网印刷设备", "position": "光伏设备龙头"},
    "688981": {"name": "中芯国际", "category": "中游制造", "segment": "晶圆代工", "main_business": "晶圆代工服务", "position": "国内最大晶圆代工厂"},
    "000725": {"name": "京东方A", "category": "中游制造", "segment": "面板/半导体", "main_business": "LCD/OLED面板", "position": "全球面板龙头"},
    "603486": {"name": "寒武纪", "category": "中游制造", "segment": "IC设计", "main_business": "AI芯片设计", "position": "AI芯片设计龙头"},
    "688223": {"name": "芯原股份", "category": "中游制造", "segment": "IC设计", "main_business": "芯片设计平台服务", "position": "芯片设计服务龙头"},
    "688521": {"name": "芯海科技", "category": "中游制造", "segment": "IC设计", "main_business": "高精度ADC、高性能MCU", "position": "模拟芯片设计"},
    "603986": {"name": "兆易创新", "category": "中游制造", "segment": "存储芯片", "main_business": "NOR Flash、MCU", "position": "国内存储芯片龙头"},
    "688008": {"name": "澜起科技", "category": "中游制造", "segment": "存储芯片", "main_business": "内存接口芯片", "position": "内存接口芯片全球龙头"},
    "002463": {"name": "长电科技", "category": "中游封装", "segment": "先进封装", "main_business": "芯片封测、CoWoS封装", "position": "全球第三大封测厂"},
    "002920": {"name": "通富微电", "category": "中游封装", "segment": "先进封装", "main_business": "集成电路封测", "position": "国内封测龙头"},
    "002185": {"name": "华天科技", "category": "中游封装", "segment": "先进封装", "main_business": "半导体封测", "position": "国内封测龙头"},
    "002916": {"name": "深南电路", "category": "中游封装", "segment": "封装载板", "main_business": "PCB、封装载板", "position": "国内PCB/载板龙头"},
    "603186": {"name": "华正新材", "category": "中游封装", "segment": "封装载板", "main_business": "覆铜板、复合材料", "position": "覆铜板材料供应商"},
    "601238": {"name": "生益科技", "category": "中游封装", "segment": "ABF薄膜", "main_business": "覆铜板、ABF薄膜", "position": "ABF薄膜上游材料"},
    "000977": {"name": "浪潮信息", "category": "下游应用", "segment": "AI服务器", "main_business": "服务器、存储", "position": "国内服务器龙头"},
    "000938": {"name": "紫光股份", "category": "下游应用", "segment": "AI服务器", "main_business": "IT基础设施、云服务", "position": "旗下新华三服务器"},
    "002837": {"name": "英维克", "category": "下游应用", "segment": "液冷散热", "main_business": "精密温控设备", "position": "数据中心温控龙头"},
    "300261": {"name": "申菱环境", "category": "下游应用", "segment": "液冷散热", "main_business": "温控设备、散热系统", "position": "数据中心温控"},
    "688169": {"name": "安克创新", "category": "下游应用", "segment": "电源管理", "main_business": "消费电子充电设备", "position": "消费电子配件龙头"},
    "002288": {"name": "*ST超华", "category": "下游应用", "segment": "PCB板", "main_business": "PCB、铜箔", "position": "PCB材料供应商"},
}


def get_stock_detail(code):
    return STOCK_INFO.get(code, {"name": code, "category": "其他", "segment": "其他", "main_business": "数据待完善", "position": "产业链位置待确定"})


def generate_docx_report(date_str, data, output_path):
    doc = Document()

    # 标题
    title = doc.add_heading(f'芯片产业链每日情报 {date_str}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'数据来源：新浪财经 | 生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

    # 大盘指数
    doc.add_heading('大盘指数', level=1)
    indices = data.get('market_indices', [])
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '指数'
    hdr[1].text = '收盘价'
    hdr[2].text = '涨跌幅'
    for idx in indices:
        row = table.add_row().cells
        row[0].text = idx.get('name', '')
        row[1].text = f"{idx.get('price', 0):.2f}"
        change = idx.get('change_pct', 0)
        row[2].text = f"{change:+.2f}%"

    # 成交量TOP15
    doc.add_heading('成交量TOP15', level=1)
    quotes = data.get('quotes', [])
    sorted_by_volume = sorted(quotes, key=lambda x: x.get('volume', 0), reverse=True)[:15]

    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '排名'
    hdr[1].text = '股票名称'
    hdr[2].text = '股票代码'
    hdr[3].text = '收盘价(元)'
    hdr[4].text = '涨跌幅'
    hdr[5].text = '成交量(万股)'
    hdr[6].text = '成交额(亿元)'

    for rank, s in enumerate(sorted_by_volume, 1):
        detail = get_stock_detail(s.get('code', ''))
        row = table.add_row().cells
        row[0].text = str(rank)
        row[1].text = detail.get('name', '')
        row[2].text = s.get('code', '')
        row[3].text = f"{s.get('price', 0):.2f}"
        change = s.get('change_pct', 0)
        row[4].text = f"{change:+.2f}%"
        # 成交量：万股 = volume/10000, 成交额：亿元 = amount/100000000
        volume = s.get('volume', 0) / 10000
        amount = s.get('amount', 0) / 100000000
        row[5].text = f"{volume:.2f}"
        row[6].text = f"{amount:.2f}"

    # 产业链详情
    categories = {"上游材料": [], "上游设备": [], "中游制造": [], "中游封装": [], "下游应用": []}
    for q in quotes:
        cat = q.get('category', '其他')
        if cat in categories:
            categories[cat].append(q)

    doc.add_heading('产业链股票详情', level=1)

    for cat_name, stocks in categories.items():
        if not stocks:
            continue
        doc.add_heading(cat_name, level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = '股票名称'
        hdr[1].text = '股票代码'
        hdr[2].text = '细分领域'
        hdr[3].text = '产业链位置'
        hdr[4].text = '涨跌幅'

        for s in stocks:
            detail = get_stock_detail(s.get('code', ''))
            row = table.add_row().cells
            row[0].text = detail.get('name', '')
            row[1].text = s.get('code', '')
            row[2].text = detail.get('segment', '')
            row[3].text = detail.get('position', '')
            change = s.get('change_pct', 0)
            row[4].text = f"{change:+.2f}%"

    doc.save(output_path)
    print(f"DOCX报告已生成: {output_path}")


if __name__ == "__main__":
    data_file = Path("data/daily_quotes/20260610/quotes.json")
    if data_file.exists():
        with open(data_file) as f:
            data = json.load(f)
        output = Path("data/supply_chain_reports/20260610/daily_digest.docx")
        output.parent.mkdir(parents=True, exist_ok=True)
        generate_docx_report("20260610", data, str(output))